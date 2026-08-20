"""Convert Markdown content into a Word document using python-docx."""
import io
import re

import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _set_base_style(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)


def _add_page_number_footer(doc: Document):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_header(doc: Document, title: str):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title or "Untitled Document")
    run.font.size = Pt(10)
    run.italic = True


def md_to_docx_bytes(markdown_text: str, title: str = "") -> bytes:
    """Render Markdown to a .docx file and return its bytes."""
    doc = Document()
    _set_base_style(doc)
    _add_header(doc, title)
    _add_page_number_footer(doc)

    if title:
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    html = markdown.markdown(markdown_text or "", extensions=["extra", "sane_lists"])
    soup = BeautifulSoup(html, "html.parser")

    for el in soup.children:
        _render_element(doc, el)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def md_to_printable_html(markdown_text: str, title: str = "") -> bytes:
    """把 Markdown 渲染为适合浏览器打印/另存为 PDF 的 HTML。"""
    body = markdown.markdown(markdown_text or "", extensions=["extra", "sane_lists", "tables"])
    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>{title or 'ResearchMate'}</title>
<style>
body {{ font-family: "Times New Roman", SimSun, serif; max-width: 820px; margin: 32px auto; line-height: 1.8; color: #1f2329; }}
h1 {{ text-align: center; font-size: 22px; }}
h2 {{ font-size: 18px; border-bottom: 1px solid #e5e7eb; padding-bottom: 4px; }}
table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
th, td {{ border: 1px solid #d1d5db; padding: 6px 10px; font-size: 14px; }}
blockquote {{ color: #6b7280; border-left: 3px solid #c7d2fe; margin: 8px 0; padding-left: 12px; }}
pre {{ background: #f3f4f6; padding: 10px; border-radius: 6px; overflow-x: auto; }}
code {{ background: #f3f4f6; padding: 2px 4px; border-radius: 4px; }}
@media print {{ body {{ margin: 12mm; }} }}
</style>
</head>
<body>
{body}
</body>
</html>
"""
    return html.encode("utf-8")


_HEADING_RE = re.compile(r"^h([1-6])$", re.I)


def _render_element(doc: Document, el):
    name = el.name.lower() if el.name else ""
    m = _HEADING_RE.match(name)
    if m:
        level = min(int(m.group(1)), 4)
        doc.add_heading(el.get_text(), level=level)
    elif name == "p":
        text = el.get_text()
        if text.strip():
            doc.add_paragraph(text)
    elif name in ("ul", "ol"):
        for li in el.find_all("li", recursive=False):
            style = "List Bullet" if name == "ul" else "List Number"
            doc.add_paragraph(li.get_text(), style=style)
    elif name == "blockquote":
        doc.add_paragraph(el.get_text(), style="Intense Quote")
    elif name == "table":
        _render_table(doc, el)
    elif name == "hr":
        doc.add_paragraph("— — — — —").alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        text = el.get_text() if el.get_text() else str(el)
        if text.strip():
            doc.add_paragraph(text)


def _render_table(doc: Document, table_el):
    rows = table_el.find_all("tr")
    if not rows:
        return
    n_cols = max(len(r.find_all(["th", "td"])) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        for j in range(n_cols):
            cell_text = cells[j].get_text() if j < len(cells) else ""
            table.cell(i, j).text = cell_text
